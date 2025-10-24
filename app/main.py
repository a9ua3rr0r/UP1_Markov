import openpyxl
from fastapi import FastAPI, HTTPException, Depends, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from pathlib import Path
from typing import List
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
import shutil, os
import random
import time
from datetime import datetime

# Импортируем только необходимые функции и классы
from app.models import (
    get_db, create_tables,
    BookCreate, BookUpdate, BookOut,
    ReaderCreate, ReaderUpdate, ReaderOut,
    BookIssueCreate, BookIssueOut,
    book_store, reader_store, book_issue_store
)

# Создаем приложение
app = FastAPI(title="LibTool", version="2.0.0")

# Пути к статическим файлам
BASE_DIR = Path(__file__).resolve().parent  # Текущая директория (app)
static_dir = BASE_DIR / 'static'
templates_dir = BASE_DIR / 'templates'

# Создаем папки если их нет
static_dir.mkdir(exist_ok=True, parents=True)
templates_dir.mkdir(exist_ok=True, parents=True)

# Монтируем статические файлы
app.mount("/static", StaticFiles(directory=static_dir), name="static")
templates = Jinja2Templates(directory=templates_dir)


# Функция очистки временных файлов сертификатов
def cleanup_temp_certificates(max_age_hours=24):
    """Очистка старых временных файлов сертификатов"""
    try:
        temp_dir = templates_dir / "temp_certificates"
        if not temp_dir.exists():
            return

        current_time = time.time()
        deleted_count = 0

        for file_path in temp_dir.glob("certificate_book_*.docx"):
            # Проверяем возраст файла
            if current_time - file_path.stat().st_mtime > max_age_hours * 3600:
                file_path.unlink()
                deleted_count += 1
                print(f"🗑️ Удален старый файл: {file_path.name}")

        if deleted_count > 0:
            print(f"✅ Очищено {deleted_count} временных файлов сертификатов")

    except Exception as e:
        print(f"⚠️ Ошибка при очистке временных файлов: {e}")


# Создаем таблицы при запуске
@app.on_event("startup")
def startup_event():
    create_tables()

    # Очищаем старые временные файлы
    cleanup_temp_certificates()

    # Автоматическая проверка просрочек при запуске
    try:
        db = next(get_db())
        updated_count = book_issue_store.check_overdue_issues(db)
        print(f"🔍 Автоматически проверены просрочки: обновлено {updated_count} выдач")
    except Exception as e:
        print(f"⚠️ Ошибка автоматической проверки просрочек: {e}")

    # Проверка директорий
    print("🔍 Проверка структуры директорий:")
    print(f"BASE_DIR: {BASE_DIR}")
    print(f"static_dir: {static_dir} (существует: {static_dir.exists()})")
    print(f"templates_dir: {templates_dir} (существует: {templates_dir.exists()})")


# Экспорт списка выдач в Excel - ОБНОВЛЕННАЯ ВЕРСИЯ
@app.get("/api/issues/export-excel")
async def export_issues_to_excel(db: Session = Depends(get_db)):
    """Экспорт списка выдач в Excel файл"""
    try:
        print("🔍 Запрос на экспорт выдач в Excel...")

        # Получаем все выдачи
        issues = book_issue_store.list_issues(db)
        print(f"📊 Найдено {len(issues)} выдач для экспорта")

        # Создаем Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Выдачи книг"

        # Стили
        header_font = Font(bold=True, size=12)
        normal_font = Font(size=10)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        center_align = Alignment(horizontal='center', vertical='center')

        # Заголовки
        headers = [
            'ID', 'Книга', 'Читатель', 'Дата выдачи',
            'Планируемый возврат', 'Фактический возврат', 'Статус'
        ]

        # Записываем заголовки
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.alignment = center_align
            cell.border = border

        # Записываем данные
        for row, issue in enumerate(issues, 2):
            # Преобразуем статус в русский текст
            status_text = {
                'issued': 'Выдана',
                'returned': 'Возвращена',
                'overdue': 'Просрочена'
            }.get(issue.status, issue.status)

            ws.cell(row=row, column=1, value=issue.id).border = border
            ws.cell(row=row, column=2, value=issue.book_name).border = border
            ws.cell(row=row, column=3, value=issue.reader_name).border = border
            ws.cell(row=row, column=4, value=issue.issue_date.strftime('%d.%m.%Y')).border = border
            ws.cell(row=row, column=5, value=issue.planned_return_date.strftime('%d.%m.%Y')).border = border

            actual_return = issue.actual_return_date.strftime('%d.%m.%Y') if issue.actual_return_date else '-'
            ws.cell(row=row, column=6, value=actual_return).border = border
            ws.cell(row=row, column=7, value=status_text).border = border

        # Настраиваем ширину колонок
        column_widths = [8, 40, 30, 12, 15, 15, 12]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # Добавляем подпись директора
        signature_row = len(issues) + 4

        # Пустая строка для разделения
        ws.cell(row=signature_row, column=1, value="")

        # Подпись директора
        signature_row += 1
        ws.merge_cells(f'A{signature_row}:G{signature_row}')
        ws.cell(row=signature_row, column=1, value="Директор библиотеки: _________________________")
        ws.cell(row=signature_row, column=1).font = Font(size=12, bold=True)
        ws.cell(row=signature_row, column=1).alignment = Alignment(horizontal='right')

        signature_row += 1
        ws.merge_cells(f'A{signature_row}:G{signature_row}')
        ws.cell(row=signature_row, column=1, value="Дата: _________________________")
        ws.cell(row=signature_row, column=1).font = Font(size=12)
        ws.cell(row=signature_row, column=1).alignment = Alignment(horizontal='right')

        # Сохраняем временный файл
        temp_dir = templates_dir / "temp_exports"
        temp_dir.mkdir(exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"выдачи_книг_{timestamp}.xlsx"
        file_path = temp_dir / filename

        wb.save(file_path)
        print(f"✅ Excel файл создан: {file_path}")
        print(f"📁 Размер файла: {file_path.stat().st_size} байт")

        # Проверяем что файл создался
        if not file_path.exists():
            raise Exception("Файл не был создан")

        # Возвращаем файл для скачивания
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    except Exception as e:
        print(f"❌ Ошибка экспорта в Excel: {str(e)}")
        import traceback
        print(f"🔍 Детали ошибки: {traceback.format_exc()}")
        raise HTTPException(500, f"Ошибка экспорта в Excel: {str(e)}")
# Главная страница
@app.get("/")
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


# API для книг
@app.get("/api/books", response_model=List[BookOut])
async def get_books(db: Session = Depends(get_db)):
    try:
        books = book_store.list_books(db)
        return books
    except Exception as e:
        raise HTTPException(500, f"Ошибка загрузки книг: {str(e)}")


@app.get("/api/books/{book_id}", response_model=BookOut)
async def get_book(book_id: int, db: Session = Depends(get_db)):
    book = book_store.get_book(db, book_id)
    if not book:
        raise HTTPException(404, "Книга не найдена")
    return book


@app.post("/api/books", response_model=BookOut)
async def create_book(book: BookCreate, db: Session = Depends(get_db)):
    try:
        return book_store.create_book(db, book)
    except Exception as e:
        raise HTTPException(500, f"Ошибка создания книги: {str(e)}")


@app.put("/api/books/{book_id}", response_model=BookOut)
async def update_book(book_id: int, book: BookUpdate, db: Session = Depends(get_db)):
    updated_book = book_store.update_book(db, book_id, book)
    if not updated_book:
        raise HTTPException(404, "Книга не найдена")
    return updated_book


@app.delete("/api/books/{book_id}")
async def delete_book(book_id: int, db: Session = Depends(get_db)):
    success = book_store.delete_book(db, book_id)
    if not success:
        raise HTTPException(404, "Книга не найдена")
    return {"ok": True}


# API для читателей
@app.get("/api/readers", response_model=List[ReaderOut])
async def get_readers(db: Session = Depends(get_db)):
    try:
        readers = reader_store.list_readers(db)
        return readers
    except Exception as e:
        raise HTTPException(500, f"Ошибка загрузки читателей: {str(e)}")


@app.post("/api/readers", response_model=ReaderOut)
async def create_reader(reader: ReaderCreate, db: Session = Depends(get_db)):
    try:
        return reader_store.create_reader(db, reader)
    except Exception as e:
        raise HTTPException(500, f"Ошибка создания читателя: {str(e)}")


@app.put("/api/readers/{reader_id}", response_model=ReaderOut)
async def update_reader(reader_id: int, reader: ReaderUpdate, db: Session = Depends(get_db)):
    updated_reader = reader_store.update_reader(db, reader_id, reader)
    if not updated_reader:
        raise HTTPException(404, "Читатель не найден")
    return updated_reader


@app.delete("/api/readers/{reader_id}")
async def delete_reader(reader_id: int, db: Session = Depends(get_db)):
    success = reader_store.delete_reader(db, reader_id)
    if not success:
        raise HTTPException(404, "Читатель не найден")
    return {"ok": True}


# API для выдачи/возврата книг
@app.get("/api/issues", response_model=List[BookIssueOut])
async def get_issues(db: Session = Depends(get_db)):
    try:
        issues = book_issue_store.list_issues(db)
        return issues
    except Exception as e:
        raise HTTPException(500, f"Ошибка загрузки выдач: {str(e)}")


@app.post("/api/issues", response_model=BookIssueOut)
async def issue_book(issue: BookIssueCreate, db: Session = Depends(get_db)):
    try:
        return book_issue_store.issue_book(db, issue)
    except Exception as e:
        raise HTTPException(500, f"Ошибка выдачи книги: {str(e)}")


@app.post("/api/issues/{issue_id}/return")
async def return_book(issue_id: int, db: Session = Depends(get_db)):
    success = book_issue_store.return_book(db, issue_id)
    if not success:
        raise HTTPException(404, "Выдача не найдена или уже возвращена")
    return {"ok": True}


# API для проверки просроченных выдач
@app.post("/api/issues/check-overdue")
async def check_overdue_issues(db: Session = Depends(get_db)):
    """Проверить и обновить статусы просроченных выдач"""
    try:
        updated_count = book_issue_store.check_overdue_issues(db)
        return {
            "updated_count": updated_count,
            "message": f"Обновлено {updated_count} просроченных выдач"
        }
    except Exception as e:
        raise HTTPException(500, f"Ошибка проверки просрочек: {str(e)}")


# API для принудительной отметки выдачи как просроченной
@app.post("/api/issues/{issue_id}/mark-overdue")
async def mark_issue_overdue(issue_id: int, db: Session = Depends(get_db)):
    """Принудительно отметить выдачу как просроченную"""
    try:
        success = book_issue_store.mark_issue_overdue(db, issue_id)
        if not success:
            raise HTTPException(400, "Не удалось отметить выдачу как просроченную")

        return {"ok": True, "message": "Выдача отмечена как просроченная"}
    except Exception as e:
        raise HTTPException(500, f"Ошибка отметки просрочки: {str(e)}")


# API для отчетов
@app.get("/api/reports/stats")
async def get_stats(db: Session = Depends(get_db)):
    try:
        # Статистика по книгам
        books = book_store.list_books(db)
        total_books = len(books)
        available_books = len([b for b in books if b.status == "available"])

        # Статистика по читателям
        readers = reader_store.list_readers(db)
        total_readers = len(readers)
        active_readers = len([r for r in readers if r.status == "active"])

        # Статистика по выдачам
        issues = book_issue_store.list_issues(db)
        total_issues = len(issues)
        current_issues = len([i for i in issues if i.status == "issued"])
        overdue_issues = len([i for i in issues if i.status == "overdue"])
        returned_issues = len([i for i in issues if i.status == "returned"])

        # Статистика по жанрам
        genre_stats = {}
        for book in books:
            genre = book.genre or "Без жанра"
            genre_stats[genre] = genre_stats.get(genre, 0) + 1

        return {
            "books": {
                "total": total_books,
                "available": available_books
            },
            "readers": {
                "total": total_readers,
                "active": active_readers
            },
            "issues": {
                "total": total_issues,
                "current": current_issues,
                "overdue": overdue_issues,
                "returned": returned_issues
            },
            "genres": genre_stats
        }
    except Exception as e:
        raise HTTPException(500, f"Ошибка загрузки статистики: {str(e)}")


# Генерация сертификата качества книги
# Упрощенная версия с фиксированной датой
@app.get("/api/certificate/{book_id}")
async def generate_certificate(book_id: int, db: Session = Depends(get_db)):
    try:
        print(f"🔍 Запрос на генерацию сертификата для книги ID: {book_id}")

        # Получаем данные книги
        book = book_store.get_book(db, book_id)
        if not book:
            raise HTTPException(404, "Книга не найдена")

        # Путь к шаблону сертификата
        template_path = templates_dir / "certificate_book_50.docx"
        if not template_path.exists():
            raise HTTPException(500, "Шаблон сертификата не найден")

        # Создаем временный файл
        cert_number = random.randint(10000, 99999)

        # Фиксируем текущую дату
        current_date = datetime.now()
        day = current_date.day
        month = current_date.month
        year = current_date.year

        # Русские названия месяцев
        month_names = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]
        month_name = month_names[month - 1]

        # Форматируем дату
        full_date = f"«{day}» {month_name} {year} г."

        temp_dir = templates_dir / "temp_certificates"
        temp_dir.mkdir(exist_ok=True)
        output_path = temp_dir / f"certificate_book_{book_id}_{cert_number}.docx"

        # Копируем шаблон
        shutil.copy(template_path, output_path)

        # Открываем документ
        doc = Document(output_path)

        # Простая замена по ключевым фразам
        for paragraph in doc.paragraphs:
            text = paragraph.text

            # Заменяем номер сертификата
            if "__________" in text and "№" in text:
                text = text.replace("__________", str(cert_number))

            # Заменяем дату
            if "«___» __________ 20___ г." in text:
                text = text.replace("«___» __________ 20___ г.", full_date)

            # Заменяем информацию о книге
            if "____________________________________________" in text and "Название:" in text:
                text = text.replace("____________________________________________", book.name)
            elif "_______________________________________________" in text and "Автор:" in text:
                text = text.replace("_______________________________________________", book.author)
            elif "________________________________________________" in text and "Жанр:" in text:
                text = text.replace("________________________________________________", book.genre or "Не указан")
            elif "______________________________" in text and "Количество экземпляров:" in text:
                text = text.replace("______________________________", str(book.count))

            # Если текст изменился, обновляем параграф
            if text != paragraph.text:
                paragraph.clear()
                paragraph.add_run(text)

        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text

                        if "_____________________" in text and "Описание повреждений" in text:
                            text = text.replace("_____________________", "отсутствуют")
                        elif "______________________________" in text and "Комментарии проверяющего" in text:
                            text = text.replace("______________________________", "книга в отличном состоянии")
                        elif "___________________________________________" in text and "Проверил:" in text:
                            text = text.replace("___________________________________________", "Иванова А.С.")
                        elif "___________________________________________" in text and "Утвердил:" in text:
                            text = text.replace("___________________________________________", "Петров И.В.")
                        elif "«___» __________ 20___ г." in text:
                            text = text.replace("«___» __________ 20___ г.", full_date)
                        elif "__________________________________________" in text and "Организация:" in text:
                            text = text.replace("__________________________________________",
                                                "Центральная городская библиотека")
                        elif "__________________________________________________" in text and "Адрес:" in text:
                            text = text.replace("__________________________________________________",
                                                "г. Москва, ул. Читательская, д. 1")
                        elif "________________________________________________" in text and "Телефон:" in text:
                            text = text.replace("________________________________________________",
                                                "+7 (495) 123-45-67")

                        if text != paragraph.text:
                            paragraph.clear()
                            paragraph.add_run(text)

        # Сохраняем изменения
        doc.save(output_path)

        return FileResponse(
            path=output_path,
            filename=f"Сертификат_качества_{book.name.replace(' ', '_')}.docx",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"❌ Ошибка генерации сертификата: {str(e)}")
        raise HTTPException(500, f"Ошибка генерации сертификата: {str(e)}")

# Скачивание правил библиотеки
@app.get("/api/rules/download")
async def download_rules():
    """Скачать правила библиотеки в формате PDF"""
    try:
        # Путь к PDF файлу с правилами
        rules_path = static_dir / "rules.pdf"

        if not rules_path.exists():
            rules_path = templates_dir / "rules.pdf"

        if not rules_path.exists():
            raise HTTPException(404, "Файл правил не найден")

        return FileResponse(
            path=rules_path,
            filename="Правила_библиотеки.pdf",
            media_type='application/pdf'
        )

    except Exception as e:
        raise HTTPException(500, f"Ошибка загрузки правил: {str(e)}")


@app.on_event("startup")
def startup_event():
    create_tables()

    # Автоматическая проверка просрочек при запуске
    try:
        updated_count = book_issue_store.check_overdue_issues(get_db())
        print(f"🔍 Автоматически проверены просрочки: обновлено {updated_count} выдач")
    except Exception as e:
        print(f"⚠️ Ошибка автоматической проверки просрочек: {e}")


# Health check
@app.get("/api/health")
async def health_check(db: Session = Depends(get_db)):
    try:
        books = book_store.list_books(db)
        return {
            "status": "healthy",
            "database": "connected",
            "books_count": len(books)
        }
    except Exception as e:
        raise HTTPException(500, f"Database error: {str(e)}")